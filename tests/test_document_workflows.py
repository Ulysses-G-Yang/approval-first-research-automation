from __future__ import annotations

import struct
import unittest
import zlib
from pathlib import Path
from tempfile import TemporaryDirectory

from research_assistant.models import TaskSpec, TaskStatus
from research_assistant.runner import TaskRunner, approve_step
from research_assistant.tools import build_default_registry
from research_assistant.workflows import available_workflows, build_workflow_plan
from research_assistant.workspace import TaskWorkspace




def _png_chunk(kind: bytes, value: bytes) -> bytes:
    return struct.pack(">I", len(value)) + kind + value + struct.pack(">I", zlib.crc32(kind + value) & 0xFFFFFFFF)


PNG_BYTES = (
    b"\x89PNG\r\n\x1a\n"
    + _png_chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    + _png_chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
    + _png_chunk(b"IEND", b"")
)


async def run_approved_plan(workspace: TaskWorkspace, plan) -> None:
    runner = TaskRunner(build_default_registry())
    for step in plan.steps:
        approve_step(workspace, step.id)
        result = await runner.resume(workspace)
        if result.task_status == TaskStatus.FAILED:
            raise AssertionError(result.message)


class DocumentWorkflowTests(unittest.IsolatedAsyncioTestCase):
    def test_bundled_workflows_are_declarative_and_versioned(self) -> None:
        self.assertTrue(
            {"research_report", "file_report", "web_to_markdown", "crawler_report", "document_to_markdown", "content_save_draft"}
            .issubset(set(available_workflows()))
        )
        with TemporaryDirectory() as temp:
            source = Path(temp) / "notes.txt"
            source.write_text("hello", encoding="utf-8")
            task = TaskSpec.create("notes", "file_report", None, [], [str(source)])
            plan = build_workflow_plan(task, build_default_registry())
        self.assertEqual([step.call.tool_name for step in plan.steps], ["file.read", "data.normalize", "report.compose"])
        self.assertEqual(plan.metadata["workflow_name"], "file_report")
        self.assertEqual(plan.metadata["definition_version"], 1)
        self.assertEqual(len(plan.metadata["definition_sha256"]), 64)

    async def test_docx_to_markdown_extracts_table_and_embedded_image(self) -> None:
        from docx import Document

        with TemporaryDirectory() as temp:
            root = Path(temp)
            image_path = root / "cover.png"
            image_path.write_bytes(PNG_BYTES)
            source = root / "article.docx"
            document = Document()
            document.add_heading("DOCX Article", level=1)
            document.add_paragraph("A paragraph for conversion.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "A"
            table.cell(1, 1).text = "1"
            document.add_picture(str(image_path))
            document.save(source)

            task = TaskSpec.create(
                "convert docx",
                "content_save_draft",
                None,
                [],
                [str(source)],
                options={"platform": "juejin"},
            )
            workspace = TaskWorkspace.create(root / "tasks", task)
            plan = build_workflow_plan(task, build_default_registry())
            workspace.save_plan(plan)
            task.status = TaskStatus.WAITING_APPROVAL
            workspace.save_task(task)
            await run_approved_plan(workspace, plan)

            markdown = workspace.read_artifact_text(workspace.list_artifacts("markdown_document")[0])
            validation = workspace.read_artifact_json(workspace.list_artifacts("markdown_validation")[0])
            draft = workspace.read_artifact_json(workspace.list_artifacts("draft_manifest")[0])
            self.assertIn("# DOCX Article", markdown)
            self.assertIn("| Name | Value |", markdown)
            self.assertIn("assets/image-001.png", markdown)
            self.assertEqual(len(workspace.list_artifacts("document_asset")), 1)
            self.assertTrue(validation["valid"])
            self.assertEqual(draft["platform"], "juejin")
            self.assertFalse(draft["network_access"])
            self.assertFalse(draft["published"])
            self.assertEqual(workspace.load_task().status, TaskStatus.COMPLETED)

    async def test_pdf_to_markdown_extracts_text_and_embedded_image(self) -> None:
        import fitz

        with TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "article.pdf"
            document = fitz.open()
            page = document.new_page()
            page.insert_text((72, 72), "PDF article body")
            page.insert_image(fitz.Rect(72, 100, 100, 128), stream=PNG_BYTES)
            document.save(source)
            document.close()

            task = TaskSpec.create("convert pdf", "document_to_markdown", None, [], [str(source)])
            workspace = TaskWorkspace.create(root / "tasks", task)
            plan = build_workflow_plan(task, build_default_registry())
            workspace.save_plan(plan)
            task.status = TaskStatus.WAITING_APPROVAL
            workspace.save_task(task)
            await run_approved_plan(workspace, plan)

            markdown = workspace.read_artifact_text(workspace.list_artifacts("markdown_document")[0])
            self.assertIn("PDF article body", markdown)
            self.assertIn("assets/image-001.png", markdown)
            self.assertEqual(len(workspace.list_artifacts("document_asset")), 1)

    async def test_markdown_input_copies_relative_images_without_network(self) -> None:
        with TemporaryDirectory() as temp:
            root = Path(temp)
            image_path = root / "cover.png"
            image_path.write_bytes(PNG_BYTES)
            source = root / "article.md"
            source.write_text("# Local Article\n\n![cover](cover.png)\n", encoding="utf-8")
            task = TaskSpec.create("convert markdown", "document_to_markdown", None, [], [str(source)])
            workspace = TaskWorkspace.create(root / "tasks", task)
            plan = build_workflow_plan(task, build_default_registry())
            workspace.save_plan(plan)
            task.status = TaskStatus.WAITING_APPROVAL
            workspace.save_task(task)
            await run_approved_plan(workspace, plan)

            markdown = workspace.read_artifact_text(workspace.list_artifacts("markdown_document")[0])
            self.assertIn("![cover](assets/image-001.png)", markdown)
            self.assertEqual(len(workspace.list_artifacts("document_asset")), 1)
